from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re

# =========================
# KONFIGURASI
# =========================
image_folder = r'C:\Users\fendi\cypress_main\cypress-main\cypress\screenshots'
output_file = 'Hasil_Pengujian_Cypress.docx'

# =========================
# FUNGSI AMBIL NOMOR TC
# =========================
def extract_tc_number(filename):
    match = re.search(r'\d+', filename)
    return int(match.group()) if match else 9999

# =========================
# BUAT DOKUMEN
# =========================
doc = Document()
doc.add_heading('Hasil Pengujian Sistem', level=1)

# =========================
# AMBIL & URUTKAN GAMBAR
# =========================
images = [
    f for f in os.listdir(image_folder)
    if f.lower().endswith(('.png', '.jpg', '.jpeg'))
]

images.sort(key=extract_tc_number)

# =========================
# MASUKKAN KE DOCX
# =========================
for img in images:
    image_path = os.path.join(image_folder, img)

    # Tambah gambar
    pic = doc.add_picture(image_path, width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Ambil nama file tanpa ekstensi
    caption_text = os.path.splitext(img)[0]

    # Tambah keterangan dari nama file
    caption = doc.add_paragraph(caption_text)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Spasi antar gambar
    doc.add_paragraph()

# =========================
# SIMPAN FILE
# =========================
doc.save(output_file)
print(f'Dokumen berhasil disimpan sebagai {output_file}')